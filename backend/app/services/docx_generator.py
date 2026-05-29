"""
Gera proposta comercial profissional em .docx usando python-docx.
"""
import base64
import io
import os
from datetime import date
from pathlib import Path
from typing import Optional
from uuid import UUID

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from app.config import get_settings
from app.models.proposal import ProposalRequest, ProposalContent

# ── Paleta de cores ──────────────────────────────────────────
BLUE_DARK  = RGBColor(0x1E, 0x3A, 0x8A)   # azul escuro (títulos)
BLUE_MID   = RGBColor(0x1D, 0x4E, 0xD8)   # azul médio (subtítulos/destaques)
BLUE_LIGHT = RGBColor(0xDB, 0xEA, 0xFE)   # azul claro (fundo cabeçalho tabela)
GRAY_TEXT  = RGBColor(0x37, 0x41, 0x51)   # cinza escuro (corpo)
GRAY_LIGHT = RGBColor(0xF1, 0xF5, 0xF9)   # cinza claro (linhas zebradas)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
TOTAL_BG   = RGBColor(0x1E, 0x3A, 0x8A)


def _currency(value: float) -> str:
    return f"R$ {value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def _set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), str(color))
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _set_cell_font(cell, bold=False, size=10, color: Optional[RGBColor] = None, italic=False):
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = bold
            run.italic = italic
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = color
        if not para.runs and para.text:
            run = para.runs[0] if para.runs else para.add_run(para.text)
            run.bold = bold
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = color


def _add_section_heading(doc: Document, text: str):
    """Heading de seção com linha azul embaixo."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE_DARK
    # borda inferior
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "1E3A8A")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_body_text(doc: Document, text: str, space_after: int = 6):
    if not text:
        return
    for line in text.split("\n"):
        p = doc.add_paragraph(line.strip())
        p.paragraph_format.space_after = Pt(space_after)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = GRAY_TEXT


def _add_info_table(doc: Document, rows: list[tuple[str, str]]):
    """Tabela de 2 colunas para dados do cliente/empresa."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    col_widths = [Cm(4), Cm(12)]
    for i, (label, value) in enumerate(rows):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        _set_cell_bg(row.cells[0], BLUE_LIGHT)
        _set_cell_font(row.cells[0], bold=True, size=9, color=BLUE_DARK)
        _set_cell_font(row.cells[1], size=10, color=GRAY_TEXT)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]


def _build_items_table(doc: Document, equipments, total: float):
    """Tabela de itens com cabeçalho azul e linhas zebradas."""
    cols = ["DESCRIÇÃO", "QTD", "VALOR UNIT.", "SUBTOTAL"]
    widths = [Cm(9), Cm(2), Cm(3.5), Cm(3.5)]
    n = len(equipments)

    table = doc.add_table(rows=1 + n + 1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # cabeçalho
    header = table.rows[0]
    for i, col in enumerate(cols):
        cell = header.cells[i]
        cell.text = col
        cell.width = widths[i]
        _set_cell_bg(cell, BLUE_DARK)
        _set_cell_font(cell, bold=True, size=9, color=WHITE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # itens
    for idx, eq in enumerate(equipments):
        subtotal = eq.quantity * eq.unit_price
        row = table.rows[idx + 1]
        values = [eq.description, str(eq.quantity), _currency(eq.unit_price), _currency(subtotal)]
        aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                  WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT]
        bg = GRAY_LIGHT if idx % 2 == 1 else WHITE
        for i, val in enumerate(values):
            cell = row.cells[i]
            cell.text = val
            cell.width = widths[i]
            cell.paragraphs[0].alignment = aligns[i]
            _set_cell_bg(cell, bg)
            _set_cell_font(cell, size=10, color=GRAY_TEXT)

    # linha de total
    total_row = table.rows[n + 1]
    total_row.cells[0].merge(total_row.cells[2])
    total_row.cells[0].text = "VALOR TOTAL DA PROPOSTA"
    total_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_cell_bg(total_row.cells[0], TOTAL_BG)
    _set_cell_font(total_row.cells[0], bold=True, size=10, color=WHITE)

    total_row.cells[3].text = _currency(total)
    total_row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_cell_bg(total_row.cells[3], TOTAL_BG)
    _set_cell_font(total_row.cells[3], bold=True, size=11, color=WHITE)


def generate_proposal_docx(
    req: ProposalRequest,
    proposal_id: UUID,
    content: Optional[ProposalContent] = None,
) -> tuple[str, float]:
    cfg = get_settings()
    storage = Path(cfg.storage_path)
    storage.mkdir(parents=True, exist_ok=True)

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── CABEÇALHO ─────────────────────────────────────────────────
    company = req.company_name or "InfraReport"

    if req.logo_base64:
        try:
            raw = req.logo_base64.split(",")[-1]
            image_stream = io.BytesIO(base64.b64decode(raw))
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_para.add_run().add_picture(image_stream, width=Inches(1.8))
        except Exception:
            pass

    p_company = doc.add_paragraph()
    p_company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_company.add_run(company)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = BLUE_DARK

    p_tagline = doc.add_paragraph()
    p_tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_tagline.add_run("PROPOSTA TÉCNICO-COMERCIAL")
    r2.font.size = Pt(12)
    r2.font.color.rgb = BLUE_MID
    r2.bold = True

    # linha divisória
    p_div = doc.add_paragraph()
    pPr = p_div._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1E3A8A")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # metadados da proposta
    meta_table = doc.add_table(rows=1, cols=3)
    meta_table.style = "Table Grid"
    meta_data = [
        ("N° Proposta", str(proposal_id)[:8].upper()),
        ("Data de Emissão", date.today().strftime("%d/%m/%Y")),
        (f"Validade", f"{req.validity_days or 30} dias"),
    ]
    for i, (label, val) in enumerate(meta_data):
        cell = meta_table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.add_run(label + "\n").bold = True
        p.runs[0].font.size = Pt(8)
        p.runs[0].font.color.rgb = BLUE_MID
        r_val = p.add_run(val)
        r_val.font.size = Pt(10)
        r_val.font.color.rgb = GRAY_TEXT
        _set_cell_bg(cell, GRAY_LIGHT)

    doc.add_paragraph()

    # ── DADOS DO CLIENTE ───────────────────────────────────────────
    _add_section_heading(doc, "Dados do Cliente")
    client_rows = [
        ("Cliente",  req.client_name),
        ("E-mail",   req.client_email),
        ("Serviço",  req.service),
    ]
    _add_info_table(doc, client_rows)

    # ── APRESENTAÇÃO ───────────────────────────────────────────────
    if content and content.introduction:
        doc.add_paragraph()
        _add_section_heading(doc, "Apresentação")
        _add_body_text(doc, content.introduction, space_after=8)

    # ── ESCOPO DO SERVIÇO ──────────────────────────────────────────
    if content and content.scope:
        doc.add_paragraph()
        _add_section_heading(doc, "Escopo do Serviço")
        _add_body_text(doc, content.scope, space_after=6)

    # ── METODOLOGIA / EXECUÇÃO ─────────────────────────────────────
    if content and content.methodology:
        doc.add_paragraph()
        _add_section_heading(doc, "Metodologia e Execução")
        _add_body_text(doc, content.methodology, space_after=6)

    # ── ITENS DA PROPOSTA ──────────────────────────────────────────
    doc.add_paragraph()
    _add_section_heading(doc, "Itens da Proposta")

    total = sum(e.quantity * e.unit_price for e in req.equipments)
    _build_items_table(doc, req.equipments, total)

    # ── CONDIÇÕES COMERCIAIS ───────────────────────────────────────
    if content and (content.payment_terms or content.execution_deadline or content.warranty):
        doc.add_paragraph()
        _add_section_heading(doc, "Condições Comerciais")

        commercial_rows = []
        if content.execution_deadline:
            commercial_rows.append(("Prazo de Execução", content.execution_deadline))
        if content.payment_terms:
            commercial_rows.append(("Condições de Pagamento", content.payment_terms))
        if content.warranty:
            commercial_rows.append(("Garantia", content.warranty))
        commercial_rows.append(("Validade da Proposta", f"{req.validity_days or 30} dias corridos"))

        _add_info_table(doc, commercial_rows)

    # ── DIFERENCIAIS ──────────────────────────────────────────────
    if content and content.differentials:
        doc.add_paragraph()
        _add_section_heading(doc, "Nossos Diferenciais")
        _add_body_text(doc, content.differentials, space_after=6)

    # ── OBSERVAÇÕES ───────────────────────────────────────────────
    notes_text = content.notes if content else req.notes
    if notes_text:
        doc.add_paragraph()
        _add_section_heading(doc, "Observações")
        _add_body_text(doc, notes_text, space_after=6)

    # ── DADOS DA EMPRESA PRESTADORA ────────────────────────────────
    doc.add_paragraph()
    _add_section_heading(doc, "Dados da Empresa")
    empresa_rows: list[tuple[str, str]] = [("Empresa", company)]
    if req.company_phone:
        empresa_rows.append(("Telefone", req.company_phone))
    if req.company_email:
        empresa_rows.append(("E-mail", req.company_email))
    if req.company_address:
        empresa_rows.append(("Endereço", req.company_address))
    _add_info_table(doc, empresa_rows)

    # ── ACEITE ────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sign = p_sign.add_run(
        f"______________________________          ______________________________\n"
        f"       {company}                               {req.client_name}\n"
        f"         (Contratante)                              (Contratado)"
    )
    r_sign.font.size = Pt(9)
    r_sign.font.color.rgb = GRAY_TEXT

    # ── RODAPÉ ────────────────────────────────────────────────────
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_f = fp.add_run(
            f"{company}  ·  Proposta Técnico-Comercial  ·  "
            f"Emitida em {date.today().strftime('%d/%m/%Y')}  ·  "
            f"Válida por {req.validity_days or 30} dias"
        )
        r_f.font.size = Pt(8)
        r_f.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    filename = f"proposta_{proposal_id}.docx"
    file_path = str(storage / filename)
    doc.save(file_path)

    return file_path, total
